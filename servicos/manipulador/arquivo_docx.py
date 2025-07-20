from servicos.manipulador.arquivo import Arquivo
from models.noticia import Noticia
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt


class ArquivoDOCX(Arquivo):
    def __init__(self):
        super().__init__()

        self.__documento = Document()



    def _formatar_titulo(self):
        titulo = self.__documento.add_heading(self.noticia.titulo, level=1)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(24)
            run.font.bold = True

    def _formatar_subtitulo(self):
        subtitulo = self.__documento.add_heading(self.noticia.subtitulo, level=2)
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in subtitulo.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(16)
            run.font.italic = True

    def _formatar_autor_data(self):
        data_formatada = (
            self.noticia.data_hora.strftime("%d/%m/%Y %H:%M:%S")
            if self.noticia.data_hora
            else "Data não informada"
        )

        autor_data = f'Por {self.noticia.autor} — {data_formatada}'



        p_meta = self.__documento.add_paragraph(autor_data)
        p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_meta.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)

    def _formatar_texto(self):
        p_texto = self.__documento.add_paragraph(self.noticia.texto)
        p_texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p_texto.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(12)
        self.__documento.add_paragraph('')

    def gerar_documento(self):
        self._formatar_titulo()
        self._formatar_subtitulo()
        self._formatar_autor_data()
        self._formatar_texto()
        self.__documento.save(self.nome_arquivo)

    def __call__(self):
        """Permite resetar o objeto chamando-o como função."""
        self.__init__()  # reinicializa tudo

